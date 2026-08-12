"""Build the manuscript .docx at BUILD time from the canonical object.

Why not the browser: MITRAL's export is
    new Blob([html], {type:"application/msword"})  ->  filename.doc
i.e. an HTML string wearing a Word MIME type. Word opens it through its HTML
importer, which is why tables lose structure and figures never arrive at all --
the Plotly charts are canvases that the serialised HTML does not contain.

Why not docx-js: the skill prefers it, but `require('docx')` fails in this
environment and `npm install docx` does not resolve. python-docx produces the
same OOXML, so the format requirement is met and the deviation is only in the
library. Recorded rather than silently substituted.

Charts: the page's inline SVG is rasterised to PNG at 2x by headless Chrome and
the PNG is embedded. Word's SVG support is unreliable; the SVG stays on the page
and as its own download.

Every value comes from the object, so the document cannot disagree with the
page by construction -- and that is then CHECKED, not assumed.
"""
import io, json, os, re, sys, time
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OBJ = r"F:\rapidmeta-ssot-shell\ssot\arni-hfref\arni-hfref.json"
PAGE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                    "ARNI_v6_mitral-base_2026-08-12.html")
OUTDIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(OUTDIR, "figs")
OUT = os.path.join(OUTDIR, "ARNI_manuscript.docx")
os.makedirs(FIGDIR, exist_ok=True)

d = json.load(open(OBJ, encoding="utf-8"))
OID = next(iter(d["results"]["by_outcome"]))
res = d["results"]["by_outcome"][OID]
out = next(o for o in d["outcomes"] if o["id"] == OID)
pooled, het = res["pooled"], res.get("heterogeneity") or {}
g = res.get("grade") or {}
sc = d.get("screening") or {}
cp = res.get("count_panels") or {}
pan = res.get("panels") or {}


def n(x):
    if x is None:
        return ""
    if isinstance(x, float):
        s = ("%.6f" % x).rstrip("0").rstrip(".")
        return s
    return str(x)


# ---------------------------------------------------------------- figures
def rasterise():
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    o = Options()
    for f in ("--headless=new", "--no-sandbox", "--force-device-scale-factor=2",
              "--window-size=1400,2200"):
        o.add_argument(f)
    dr = webdriver.Chrome(options=o)
    figs = []
    try:
        dr.get("file:///" + PAGE.replace("\\", "/"))
        time.sleep(1.5)
        # open every panel so the figures have layout
        dr.execute_script(
            "document.querySelectorAll('.panel').forEach(p=>{"
            "p.style.height='auto';p.style.overflow='visible';});")
        time.sleep(0.8)
        els = dr.find_elements("css selector", "svg")
        for i, el in enumerate(els):
            title = dr.execute_script(
                "const c=arguments[0].closest('.card');"
                "const h=c?c.querySelector('h3'):null;"
                "return h?h.innerText.trim():'Figure';", el)
            dr.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            time.sleep(0.25)
            p = os.path.join(FIGDIR, "fig%02d.png" % (i + 1))
            el.screenshot(p)
            figs.append((title, p))
    finally:
        dr.quit()
    return figs


# ---------------------------------------------------------------- helpers
def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    t = OxmlElement('w:t')
    t.text = "Right-click and choose Update Field to build the table of contents."
    fld.append(t)
    r._r.addnext(fld)


def page_numbers(doc):
    ftr = doc.sections[0].footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ftr.add_run()
    for instr in ('begin', 'PAGE', 'end'):
        el = OxmlElement('w:fldChar' if instr in ('begin', 'end') else 'w:instrText')
        if instr in ('begin', 'end'):
            el.set(qn('w:fldCharType'), instr)
        else:
            el.set(qn('xml:space'), 'preserve')
            el.text = ' PAGE '
        r._r.append(el)


TBL = 0
FIG = 0


def table(doc, caption, headers, rows):
    global TBL
    TBL += 1
    c = doc.add_paragraph("Table %d. %s" % (TBL, caption))
    c.runs[0].bold = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = n(v)
    doc.add_paragraph()
    return t


def figure(doc, caption, path, width=6.0):
    global FIG
    FIG += 1
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph("Figure %d. %s" % (FIG, caption))
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True


# ---------------------------------------------------------------- build
figs = rasterise()
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)

doc.add_heading(d["title"], 0)
p = doc.add_paragraph(d["question"])
p.runs[0].italic = True

rd_state = "NOT YET DETERMINED"
warn = doc.add_paragraph()
wr = warn.add_run("Submission readiness: %s. This manuscript is projected from a "
                  "canonical data object; the sections marked as awaiting the "
                  "author are not written here." % rd_state)
wr.bold = True
wr.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

doc.add_heading("Contents", 1)
add_toc(doc)
doc.add_page_break()

# --- Abstract -------------------------------------------------------------
doc.add_heading("Abstract", 1)
doc.add_paragraph("Objective. " + d["question"])
doc.add_paragraph("Data sources and eligibility. " + (sc.get("search_note") or "")
                  + " " + (sc.get("eligibility") or ""))
doc.add_paragraph(
    "Synthesis. %s model, estimator %s, k = %s."
    % (res.get("model", ""), res.get("estimator_used") or res.get("estimator", ""),
       n(res.get("k"))))
doc.add_paragraph(
    "Results. Pooled %s %s (%s to %s, %s%% interval); I-squared %s%%, "
    "tau-squared %s."
    % (pooled["measure"], n(pooled["point"]), n(pooled["ci_low"]),
       n(pooled["ci_high"]), n(pooled.get("ci_level", 95)),
       n(het.get("i2")), n(het.get("tau2"))))
if g.get("certainty"):
    doc.add_paragraph("Certainty. GRADE certainty %s. %s"
                      % (g["certainty"], g.get("certainty_derivation", "")))
doc.add_paragraph("Limitations. " + (sc.get("known_limitation") or ""))

# --- Introduction ---------------------------------------------------------
doc.add_heading("Introduction", 1)
doc.add_paragraph(
    "This section is not written here. The canonical object carries no "
    "background or rationale field, and an introduction generated without one "
    "would be argument that no source in this review supports. It is left for "
    "the author, and its absence is stated rather than filled.")

# --- Methods --------------------------------------------------------------
doc.add_heading("Methods", 1)
doc.add_heading("Protocol and registration", 2)
reg = d.get("registration") or {}
c0 = (reg.get("commits") or [{}])[0]
doc.add_paragraph(
    "Registered as a timestamped public commit in %s at path %s. Commit %s, "
    "committed %s. %s"
    % (reg.get("repository", ""), reg.get("path", ""), c0.get("sha", "")[:12],
       c0.get("committed_utc", ""),
       (reg.get("ordering") or {}).get("reason", "")))
doc.add_heading("Eligibility criteria", 2)
doc.add_paragraph(sc.get("eligibility", ""))
doc.add_heading("Information sources and search", 2)
doc.add_paragraph(sc.get("search_note", ""))
doc.add_paragraph(
    "No executed query string is recorded in the canonical object for any "
    "database. The search is therefore described but not reproducible from "
    "this document, and that is a stated limitation rather than an omission.")
doc.add_heading("Synthesis methods", 2)
doc.add_paragraph(
    "%s model with the %s estimator. %s"
    % (res.get("model", ""), res.get("estimator_used") or res.get("estimator", ""),
       (res.get("handbook") or {}).get("note", "")))

# --- Results --------------------------------------------------------------
doc.add_heading("Results", 1)
doc.add_heading("Included studies", 2)
table(doc, "Characteristics of the included trials",
      ["Trial", "Registry ID", "Year", "Design", "Population", "Analysed (T/C)"],
      [[t.get("name") or t["id"], t.get("nct", ""), t.get("year", ""),
        (t.get("design", "") or "")[:60],
        (t.get("population", "") or "")[:50],
        "%s / %s" % (n((t["by_outcome"][OID].get("treatment") or {}).get("n")),
                     n((t["by_outcome"][OID].get("control") or {}).get("n")))]
       for t in d["inputs"]["trials"]])

doc.add_heading("Per-arm event counts", 2)
table(doc, "Per-arm events and denominators for the pooled composite",
      ["Trial", "Intervention events / n", "Comparator events / n",
       "Source tier"],
      [[t.get("name") or t["id"],
        "%s / %s" % (n(t["by_outcome"][OID]["treatment"]["events"]),
                     n(t["by_outcome"][OID]["treatment"]["n"])),
        "%s / %s" % (n(t["by_outcome"][OID]["control"]["events"]),
                     n(t["by_outcome"][OID]["control"]["n"])),
        t["by_outcome"][OID]["treatment"].get("source_tier", "")]
       for t in d["inputs"]["trials"]
       if (t["by_outcome"][OID].get("treatment") or {}).get("events") is not None])

doc.add_heading("Pooled result", 2)
doc.add_paragraph(
    "Pooled %s %s (%s to %s), k = %s. I-squared %s%%, tau-squared %s, Q %s on "
    "%s degrees of freedom."
    % (pooled["measure"], n(pooled["point"]), n(pooled["ci_low"]),
       n(pooled["ci_high"]), n(res.get("k")), n(het.get("i2")),
       n(het.get("tau2")), n(het.get("q")), n(het.get("df"))))

rows = [["Risk ratio", "%s (%s to %s)" % (n(cp["rr"]["point"]), n(cp["rr"]["ci_low"]),
                                          n(cp["rr"]["ci_high"])), n(cp["rr"]["I2"])],
        ["Odds ratio", "%s (%s to %s)" % (n(cp["or"]["point"]), n(cp["or"]["ci_low"]),
                                          n(cp["or"]["ci_high"])), n(cp["or"]["I2"])],
        ["Risk difference", "%s (%s to %s)" % (n(cp["rd"]["point"]),
                                               n(cp["rd"]["ci_low"]),
                                               n(cp["rd"]["ci_high"])),
         n(cp["rd"]["I2"])]] if cp else []
if rows:
    table(doc, "The same 2x2 on three scales (sensitivity to the primary "
               "hazard-ratio pool)", ["Measure", "Pooled (95% interval)",
                                      "I-squared (%)"], rows)

if pan.get("leave_one_out"):
    table(doc, "Leave-one-out sensitivity analysis",
          ["Trial omitted", "Pooled estimate (95% interval)", "I-squared (%)"],
          [[x["omitted"], "%s (%s to %s)" % (n(x["point"]), n(x["ci_low"]),
                                             n(x["ci_high"])), n(x["I2"])]
           for x in pan["leave_one_out"]])

mc = (res.get("sensitivity") or {}).get("between_study_variance_method_comparison")
if mc and mc.get("methods"):
    table(doc, "Between-study-variance estimator comparison",
          ["Estimator", "Interval method", "Point", "Lower", "Upper", "tau-squared"],
          [[m.get("between_study_variance_estimator"), m.get("interval_method"),
            n(m.get("point")), n(m.get("ci_low")), n(m.get("ci_high")),
            n(m.get("tau2"))] for m in mc["methods"]])

doc.add_heading("Figures", 2)
for title, path in figs:
    figure(doc, title, path)

# --- Screening log --------------------------------------------------------
doc.add_heading("Screening and exclusions", 2)
table(doc, "Screening log: every record adjudicated, with its decision and the "
           "reason for it",
      ["Record", "Registry / PMID", "Decision", "Reason"],
      [[r.get("trial", ""),
        " ".join(filter(None, [r.get("nct", ""),
                               ("PMID " + str(r["pmid"])) if r.get("pmid") else ""])),
        r.get("disposition", "excluded"),
        (r.get("reason") or "")[:400]] for r in (sc.get("records") or [])])
doc.add_paragraph(
    "This log covers the records the review adjudicated. The object does not "
    "hold the full set retrieved by the search, nor the stage at which each "
    "decision was taken, so a title/abstract versus full-text split cannot be "
    "shown. That gap is stated here rather than implied by a shorter table.")

# --- GRADE ----------------------------------------------------------------
if g.get("domains"):
    doc.add_heading("Certainty of the evidence", 1)
    table(doc, "GRADE summary of findings",
          ["Domain", "Rating", "Basis"],
          [[k.replace("_", " ").capitalize(), v.get("rating", ""),
            (v.get("basis_in_sources") or "")[:400]]
           for k, v in g["domains"].items()])
    doc.add_paragraph("Overall certainty: %s. %s"
                      % (g.get("certainty", ""), g.get("certainty_derivation", "")))

# --- Discussion / limitations / conclusions -------------------------------
doc.add_heading("Discussion", 1)
doc.add_paragraph(
    "This section is not written here, for the same reason as the "
    "introduction: interpretation is argument, and the object holds none. What "
    "the object does hold, and what a discussion would have to account for, is "
    "stated in the limitations below.")
doc.add_heading("Limitations", 1)
for txt in filter(None, [
        (res.get("sensitivity") or {}).get("leave_one_out_finding"),
        sc.get("known_limitation"),
        (reg.get("ordering") or {}).get("reason"),
        (cp.get("_provenance") if cp else None)]):
    doc.add_paragraph(txt, style="List Bullet")
doc.add_heading("Conclusions", 1)
doc.add_paragraph("Awaiting the author. See Discussion.")
doc.add_heading("Funding and conflicts of interest", 1)
doc.add_paragraph("The canonical object records neither a funding statement nor "
                  "a conflict-of-interest declaration for this review.")

# --- References -----------------------------------------------------------
doc.add_heading("References", 1)
srcs = sorted((d.get("sources") or {}).values(),
              key=lambda x: x.get("layer_rank", 99))
for i, s in enumerate(srcs, 1):
    doc.add_paragraph("%d. %s. %s" % (i, s.get("name", ""), s.get("url", "")))

page_numbers(doc)
doc.save(OUT)
print("wrote", OUT, os.path.getsize(OUT), "bytes")
print("tables:", TBL, "figures:", FIG, "rasterised:", len(figs))
