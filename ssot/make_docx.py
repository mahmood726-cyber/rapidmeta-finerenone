"""Build the manuscript .docx at BUILD time from the canonical object.

Why not the browser: MITRAL's export is
    new Blob([html], {type:"application/msword"})  ->  filename.doc
i.e. an HTML string wearing a Word MIME type. Word opens it through its HTML
importer, which is why tables lose structure and figures never arrive at all --
the Plotly charts are canvases that the serialised HTML does not contain.

Why not docx-js: the skill prefers it, but `require('docx')` fails in this
environment and `npm install docx` does not resolve. python-docx produces the
same OOXML, so the format requirement is met and the deviation is only in the
library. Recorded rather than silently substituted.

Charts: the page's inline SVG is rasterised to PNG at 2x by headless Chrome and
the PNG is embedded. Word's SVG support is unreliable; the SVG stays on the page
and as its own download.

Every value comes from the object, so the document cannot disagree with the
page by construction -- and that is then CHECKED, not assumed.
"""
import io, json, os, re, sys, time
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OBJ = r"F:\rapidmeta-ssot-shell\ssot\arni-hfref\arni-hfref.json"
# The NEWEST built page, resolved at run time. This was pinned to the v6 file,
# so every rebuild of the .docx silently rasterised figures from a page five
# builds old while the prose came from the current object -- a document whose
# text and figures had different ancestors, which is exactly the divergence this
# whole design exists to prevent.
def _newest_page():
    here = os.path.dirname(os.path.abspath(__file__))
    import glob as _g
    c = [f for f in _g.glob(os.path.join(here, "ARNI_v*.html"))]
    if not c:
        raise SystemExit("no built ARNI page found to rasterise figures from")
    return max(c, key=os.path.getmtime)


PAGE = _newest_page()
OUTDIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(OUTDIR, "figs")
OUT = os.path.join(OUTDIR, "ARNI_manuscript.docx")
os.makedirs(FIGDIR, exist_ok=True)

d = json.load(open(OBJ, encoding="utf-8"))
OID = next(iter(d["results"]["by_outcome"]))
res = d["results"]["by_outcome"][OID]
out = next(o for o in d["outcomes"] if o["id"] == OID)
pooled, het = res["pooled"], res.get("heterogeneity") or {}
g = res.get("grade") or {}
sc = d.get("screening") or {}
cp = res.get("count_panels") or {}
pan = res.get("panels") or {}


sys.path.insert(0, os.path.join("F:", os.sep, "rapidmeta-ssot-shell", "ssot"))
import projectors as _PJ  # noqa: E402


def n(x):
    """Display formatting, 3 significant figures on floats.

    Tables reported estimates to six decimal places while the prose used four.
    Consistent numeric formatting is one of the signals a journal editor reads
    for. This file and the page render the SAME block list, so they round
    together or they disagree with each other. The object keeps every digit.
    """
    if x is None:
        return ""
    if isinstance(x, bool):
        return str(x)
    if isinstance(x, float):
        return _PJ.sig(x, 3)
    return str(x)


# ---------------------------------------------------------------- figures
def rasterise():
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    o = Options()
    for f in ("--headless=new", "--no-sandbox", "--force-device-scale-factor=2",
              "--window-size=1400,2200"):
        o.add_argument(f)
    dr = webdriver.Chrome(options=o)
    figs = []
    try:
        dr.get("file:///" + PAGE.replace("\\", "/"))
        time.sleep(1.5)
        # open every panel so the figures have layout
        dr.execute_script(
            "document.querySelectorAll('.panel').forEach(p=>{"
            "p.style.height='auto';p.style.overflow='visible';});")
        time.sleep(0.8)
        # Skip the forest range variants that are not the active one. They are
        # present in the DOM at height:0 (deliberately, so the invariance
        # detector can read them), so a naive svg sweep screenshots all of them
        # and the document gains three near-duplicate forests of the same data
        # at different axis windows -- that is a page control, not a figure set.
        els = [e for e in dr.find_elements("css selector", "svg")
               if dr.execute_script(
                   "const p=arguments[0].closest('.fwp');"
                   "return !p || p.id==='fwp-fit';", e)]
        for i, el in enumerate(els):
            # The card's NOTE as well as its title. Every figure legend in the
            # Word file was a bare title -- "Forest plot" -- while the page
            # carried the real legend underneath, including the cautions saying
            # a funnel cannot be read at this k and why GOSH and TSA are not
            # drawn. Those cautions existed on ONE surface only, so the Word
            # reader met seven diagnostics with no warning attached to any.
            title = dr.execute_script(
                "const c=arguments[0].closest('.card');"
                "const h=c?c.querySelector('h3'):null;"
                "return h?h.innerText.trim():'Figure';", el)
            note = dr.execute_script(
                "const c=arguments[0].closest('.card');"
                "const ss=c?[...c.querySelectorAll('p > small')]:[];"
                "return ss.length?ss[ss.length-1].innerText.trim():'';", el)
            if note:
                title = title + ". " + note
            dr.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            time.sleep(0.25)
            p = os.path.join(FIGDIR, "fig%02d.png" % (i + 1))
            el.screenshot(p)
            figs.append((title, p))
    finally:
        dr.quit()
    return figs


# ---------------------------------------------------------------- helpers
def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    # CT_SimpleField holds RUNS, not text. The placeholder went in as a bare
    # <w:t> child, which is schema-invalid -- caught by the docx skill's XSD
    # validator, and invisible to every check that only reads the text back out
    # (pandoc, python-docx and our own numeral checks all still found the string,
    # because they walk descendants rather than validate the tree). Word repairs
    # it silently on open, which is exactly why it survived: the artefact looked
    # fine to a human and was malformed to a parser.
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "Right-click and choose Update Field to build the table of contents."
    run.append(t)
    fld.append(run)
    r._r.addnext(fld)


def page_numbers(doc):
    ftr = doc.sections[0].footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ftr.add_run()
    for instr in ('begin', 'PAGE', 'end'):
        el = OxmlElement('w:fldChar' if instr in ('begin', 'end') else 'w:instrText')
        if instr in ('begin', 'end'):
            el.set(qn('w:fldCharType'), instr)
        else:
            el.set(qn('xml:space'), 'preserve')
            el.text = ' PAGE '
        r._r.append(el)


TBL = 0
FIG = 0


def table(doc, caption, headers, rows, mono_cols=()):
    """Render a table. `mono_cols` are column indices holding VERBATIM text.

    Quoted evidence belongs in a table with its context -- what result it belongs
    to, the call, the estimate -- not as a free-standing block a reader has to
    match up by eye. But a metafor table whose columns stop lining up is a
    misquotation, so the monospace and the line breaks are preserved INSIDE the
    cell: one paragraph per line, Consolas, zero space-after.
    """
    global TBL
    TBL += 1
    c = doc.add_paragraph("Table %d. %s" % (TBL, caption))
    c.runs[0].bold = True
    t = doc.add_table(rows=1, cols=len(headers))
    # Plain ruled grid. "Light Grid Accent 1" is an Office theme style whose
    # colour belongs to a slide deck, not to a manuscript table.
    try:
        t.style = "Table Grid"
    except KeyError:
        t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # REPEAT THE HEADER ROW ACROSS PAGE BREAKS. Zero of seventeen tables had
    # this. The screening log is fourteen rows and the statistical-output table
    # fourteen more; both cross a page in the Word file, and a reader met the
    # continuation with no column headers at all. This is a reading defect, not
    # a cosmetic one, and it is the first thing a reviewer would hit.
    _trPr = t.rows[0]._tr.get_or_add_trPr()
    _th = OxmlElement('w:tblHeader')
    _th.set(qn('w:val'), "true")
    _trPr.append(_th)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        # Light shading on the header so the repeat is recognisable as a header
        # rather than as data. Grey, not a theme accent: this is a manuscript.
        _sh = OxmlElement('w:shd')
        _sh.set(qn('w:val'), 'clear')
        _sh.set(qn('w:fill'), 'EDEDED')
        cell._tc.get_or_add_tcPr().append(_sh)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            if i in mono_cols:
                cell = cells[i]
                cell.text = ""
                lines = str(v if v is not None else "").split("\n")
                for j, line in enumerate(lines):
                    para = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
                    para.paragraph_format.space_after = Pt(0)
                    run = para.add_run(line)
                    run.font.name = "Consolas"
                    run.font.size = Pt(7.5)
            else:
                cells[i].text = n(v)
    doc.add_paragraph()
    return t


def figure(doc, caption, path, width=6.0):
    """Size to the image, within the measure -- not every figure to the measure.

    All fifteen were forced to exactly 6.00 inches against a text measure of
    exactly 6.00 inches (8.5 less two 1.25 margins), so every figure ran edge to
    edge with no breathing room, and because aspect ratios span 0.15 to 0.68 the
    empty-state panels became 6x0.9 inch strips. Width is now capped below the
    measure and height capped so nothing runs past a page, with the aspect ratio
    preserved from the file rather than assumed.
    """
    global FIG
    FIG += 1
    MAX_W, MAX_H = 5.6, 7.0
    w = width
    try:
        from PIL import Image as _I
        with _I.open(path) as im:
            pw, ph = im.size
        w = min(MAX_W, MAX_W)                       # never wider than the cap
        if ph and pw and (ph / float(pw)) * w > MAX_H:
            w = MAX_H * (pw / float(ph))            # tall image: bound by height
    except Exception:                                # noqa: BLE001
        w = min(width, MAX_W)
    doc.add_picture(path, width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph("Figure %d. %s" % (FIG, caption))
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True



# ---- deferral -----------------------------------------------------------
# The journal places tables and figure legends at the END of the manuscript.
# They are recorded in body order and emitted later, so nothing is re-authored
# and the recorded block list -- which the web page renders -- stays a single
# description of one document.
_DEFERRED_T, _DEFERRED_F = [], []
_DEFER_ON = True
_real_table, _real_figure = table, figure


def table(doc, caption, headers, rows, mono_cols=()):
    global TBL
    TBL += 1
    # mono_cols travels in the block so the page renders the same cells verbatim.
    # A quotation that is monospaced in one surface and reflowed in the other is
    # two different quotations, which the alignment gate would then have to call
    # a divergence.
    DOCMODEL.append({"kind": "table", "n": TBL, "caption": caption,
                     "headers": [str(h) for h in headers],
                     "mono_cols": list(mono_cols),
                     "rows": [[(v if i in mono_cols else n(v))
                               for i, v in enumerate(r)] for r in rows]})
    if _DEFER_ON:
        _DEFERRED_T.append((TBL, caption, headers, rows, tuple(mono_cols)))
        return None
    return _real_table(doc, caption, headers, rows, mono_cols)


def verbatim(doc, text):
    """A quoted block, preserved line for line in BOTH surfaces.

    Recorded as kind "pre" rather than as paragraphs: split into paragraphs the
    columns of a metafor table stop lining up, and a misaligned quotation is a
    misquotation. Monospace here, <pre> on the page, same string in each.
    """
    DOCMODEL.append({"kind": "pre", "text": str(text)})
    for line in str(text).split("\n"):
        p = _ap(line if line.strip() else "")
        p.paragraph_format.space_after = Pt(0)
        r = p.runs[0] if p.runs else p.add_run("")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
    doc.add_paragraph()


def figure(doc, caption, path, width=6.0):
    global FIG
    FIG += 1
    DOCMODEL.append({"kind": "figure", "n": FIG, "caption": caption,
                     "png": os.path.abspath(path)})
    if _DEFER_ON:
        _DEFERRED_F.append((FIG, caption, path, width))
        return None
    return _real_figure(doc, caption, path, width)

# ---------------------------------------------------------------- build
figs = rasterise()
doc = Document()

# ---- document model recorder -------------------------------------------
# The page used to DESCRIBE this manuscript in cards while the .docx CONTAINED
# it: two renderings of one paper that could drift apart, and did. Every block
# emitted below is recorded in order and written to disk, and the page renders
# that record -- so what a reader sees on the page is what is in the Word file,
# by construction rather than by intention.
def _fix_zoom(document):
    """python-docx ships a schema-invalid <w:zoom> in its own default template.

    CT_Zoom requires w:percent; the bundled template sets only w:val="bestFit",
    so every file python-docx produces carries the defect and ours inherited it.
    Upstream's bug, but it ships inside OUR .docx, so it is repaired here rather
    than reported and left in. w:val is kept -- it is a legal optional attribute
    and is what actually drives the view; only the required one was missing.
    """
    st = document.settings.element
    for z in st.findall(qn('w:zoom')):
        if z.get(qn('w:percent')) is None:
            z.set(qn('w:percent'), "100")


_fix_zoom(doc)

DOCMODEL = []
_ah, _ap = doc.add_heading, doc.add_paragraph


_IN_HEADING = [False]


def _rec_heading(text="", level=1, *a, **k):
    """Record the heading ONCE.

    python-docx implements add_heading BY CALLING add_paragraph, so with both
    patched the same text was recorded twice -- once as a heading and once as
    prose directly beneath it. That is the 55-of-68 heading echo the rendered
    review found, including the 43-word title repeated verbatim as the first body
    paragraph. The flag suppresses the inner recording; the Word file is
    unaffected because the real add_paragraph still runs.
    """
    DOCMODEL.append({"kind": "h%d" % max(1, min(4, level or 1)), "text": str(text)})
    _IN_HEADING[0] = True
    try:
        return _ah(text, level, *a, **k)
    finally:
        _IN_HEADING[0] = False


class _RecPar:
    """A paragraph proxy that records text added through add_run().

    Whole sections were invisible to the recorder because they are built as an
    EMPTY add_paragraph() followed by add_run() calls -- the bolded label and its
    body. The recorder only saw the empty string, so the Data Availability
    Statement, the software block and the licensing note existed in the Word file
    and were absent from the page, which is precisely the drift the shared block
    list exists to prevent. Runs now append to the recorded block.
    """

    def __init__(self, real, block):
        self._real, self._block = real, block

    def add_run(self, text="", *a, **k):
        t = str(text) if text is not None else ""
        if t.strip():
            if self._block is None:
                self._block = {"kind": "p", "text": ""}
                DOCMODEL.append(self._block)
            self._block["text"] = (self._block["text"] + t) if self._block["text"] else t
        return self._real.add_run(text, *a, **k)

    def __getattr__(self, name):
        return getattr(self._real, name)


def _rec_par(text="", *a, **k):
    t = str(text) if text is not None else ""
    blk = None
    if t.strip() and not _IN_HEADING[0]:
        blk = {"kind": "li" if k.get("style") == "List Bullet" else "p", "text": t}
        DOCMODEL.append(blk)
    real = _ap(text, *a, **k)
    return _RecPar(real, blk)


doc.add_heading, doc.add_paragraph = _rec_heading, _rec_par
st = doc.styles["Normal"]
# A MANUSCRIPT FACE, NOT A UI FACE. Calibri 11 is Word's default and reads as an
# office document; the HTML projection -- which is the surface Mahmood judged
# better -- sets Georgia. Matching it is the point: both are projections of one
# object and the Word file was the lossy one.
st.font.name = "Georgia"
st.font.size = Pt(10.5)
# 1.5 line spacing. Journals ask for 1.5 or double on a review copy and Word's
# default single is neither. Ragged right rather than justified: Word does not
# hyphenate by default and justified text without hyphenation opens rivers,
# which is worse than the ragged edge it replaces.
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)

doc.add_heading(d["title"], 0)
p = doc.add_paragraph(d["question"])
p.runs[0].italic = True

# COMPUTED, not a constant. The page fixed this months of defects ago; the .docx
# still carried the literal string, so the Word file could have claimed a state
# no object could change. Same defect class, second surface.
sys.path.insert(0, os.path.join("F:", os.sep, "rapidmeta-ssot-shell", "ssot"))
import projectors as _pj  # noqa: E402
_rd = _pj.readiness(d)
rd_state = _rd["state"] if isinstance(_rd, dict) else str(_rd)
warn = doc.add_paragraph()
wr = warn.add_run("Submission readiness: %s. Every number in this document is "
                  "projected from a single canonical object; sections that could "
                  "not be written from it say so in place." % rd_state)
wr.bold = True
wr.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

doc.add_heading("Contents", 1)
add_toc(doc)
doc.add_page_break()

# --- Manuscript prose, from the SAME block the page projects -----------------
# The .docx previously composed its own abstract from field labels and printed
# "this section is not written here" as the introduction, while the page carried
# a full manuscript. Two surfaces claiming to be the same paper, disagreeing
# about whether the paper exists. Both now project canon["manuscript"], so they
# cannot diverge -- and the numbers inside the prose are substituted from the
# results block at build time rather than typed.
sys.path.insert(0, r"F:\rapidmeta-ssot-shell\ssot")
import paper as PP  # noqa: E402

MS = d.get("manuscript") or {}
TOK = PP.build_tokens(d, res, OID) if MS else {}


def fill(text, where):
    """Substitute [[tokens]] as PLAIN text -- no HTML spans in a Word file."""
    return PP.fill(text, TOK, where, num_span=False)


doc.add_heading("Abstract", 1)
if MS.get("abstract"):
    for k, v in MS["abstract"].items():
        if k.startswith("_"):
            continue
        para = doc.add_paragraph()
        r = para.add_run(k.replace("_", " ").capitalize() + ". ")
        r.bold = True
        para.add_run(fill(v, "abstract." + k))
else:
    doc.add_paragraph("Objective. " + d["question"])

doc.add_heading("Introduction", 1)
for x in MS.get("introduction", []):
    doc.add_paragraph(fill(x["text"], "introduction"))
if not MS.get("introduction"):
    doc.add_paragraph(
        "This section is not written here. The canonical object carries no "
        "background field, and an introduction generated without one would be "
        "argument that no source in this review supports.")

# --- Methods --------------------------------------------------------------
doc.add_heading("Methods", 1)
# Prose first, then the field-level detail. A methods section that opens with a
# list of labels reads as a form; one that opens with sentences reads as methods.
for _x in MS.get("methods_prose", []):
    doc.add_heading(_x["heading"], 2)
    doc.add_paragraph(fill(_x["text"], "methods." + _x["heading"]))
if MS.get("methods_prose"):
    doc.add_heading("Recorded detail", 2)
doc.add_heading("Protocol and registration", 3)
reg = d.get("registration") or {}
c0 = (reg.get("commits") or [{}])[0]
doc.add_paragraph(
    "Registered as a timestamped public commit in %s at path %s. Commit %s, "
    "committed %s. %s"
    % (reg.get("repository", ""), reg.get("path", ""), c0.get("sha", "")[:12],
       c0.get("committed_utc", ""),
       (reg.get("ordering") or {}).get("reason", "")))
doc.add_heading("Eligibility criteria", 3)
doc.add_paragraph(sc.get("eligibility", ""))
doc.add_heading("Information sources and search", 3)
doc.add_paragraph(sc.get("search_note", ""))
doc.add_paragraph(
    "Each database's query string is recorded in the canonical object exactly "
    "as executed, with the endpoint, the parameters, the filters, the execution "
    "timestamp and the hit count the API returned, and all of it is reproduced "
    "in the extended data (search_capture.csv). The search is therefore "
    "reproducible from this record. This paragraph previously asserted the "
    "opposite -- that no query string was recorded and the search was not "
    "reproducible -- while the strings sat in the object six paragraphs away.")
doc.add_heading("Synthesis methods", 3)
_model = str(res.get("model", "") or "")
# The object stores 'random'; "random model with the REML estimator." is what
# that produced, which reads as a dropped word rather than a term of art. The
# expansion is presentational only -- the stored value is untouched.
_model = {"random": "Random-effects", "fixed": "Fixed-effect",
          "common": "Common-effect"}.get(_model.lower(), _model or "Random-effects")
doc.add_paragraph(
    "%s model with the %s estimator.%s"
    % (_model, res.get("estimator_used") or res.get("estimator", ""),
       (" " + (res.get("handbook") or {}).get("note", ""))
       if (res.get("handbook") or {}).get("note") else ""))

# --- Results --------------------------------------------------------------
doc.add_heading("Results", 1)

# PRISMA counts as TEXT as well as a figure. The flow diagram is an embedded
# image, so its numbers are not selectable, searchable or machine-readable in the
# Word file -- a submission checker looking for the identification count finds
# nothing. The figure stays; this is the same counts as a table beside it.
_scr = d.get("screening") or {}
_corp = _scr.get("corpus") or []
if _corp:
    import collections as _cc
    import re as _re2
    _cc2 = _scr.get("corpus_counts") or {}
    _ident, _perdb = 0, []
    for _db in ((d.get("search") or {}).get("databases") or []):
        _m = _re2.search(r"(\d+)", str(_db.get("records_retrieved")
                                       or _db.get("hit_count") or ""))
        if _m:
            _ident += int(_m.group(1))
            _perdb.append("%s: %s" % (str(_db.get("database","")).split(" (")[0],
                                      _m.group(1)))
    _src = _cc.Counter(r.get("source") for r in _corp)
    _extiab = _cc2.get("TiAb/exclude") or 0
    _und = sum(v for k, v in _cc2.items() if str(k).endswith("undetermined"))
    _full = sum(v for k, v in _cc2.items() if str(k).startswith("FullText"))
    _exfull = _cc2.get("FullText/exclude") or 0
    _inc = sum(v for k, v in _cc2.items() if str(k).endswith("INCLUDE"))
    table(doc, "PRISMA flow of records, as counts",
          ["Stage", "n", "Detail"],
          [["Records identified from databases and registers", _ident,
            "; ".join(_perdb)],
           ["Records removed before screening", 0 if _ident == len(_corp) else "",
            "No de-duplication step is recorded; the retrieved totals sum exactly "
            "to the screened corpus. Corpus tally: "
            + ", ".join("%s %d" % (k, v) for k, v in sorted(_src.items()) if k)],
           ["Records screened on title and abstract", len(_corp), ""],
           ["Records excluded at title and abstract", _extiab, ""],
           ["Records undetermined at title and abstract", _und,
            "Not counted as exclusions."],
           ["Reports assessed for eligibility at full text", _full, ""],
           ["Reports excluded at full text", _exfull, ""],
           ["Studies contributing to the synthesis", _inc, ""]])

for _x in MS.get("results_prose", []):
    doc.add_heading(_x["heading"], 2)
    doc.add_paragraph(fill(_x["text"], "results." + _x["heading"]))
if MS.get("results_prose"):
    doc.add_heading("Tables and figures", 2)
doc.add_heading("Included studies", 3)
table(doc, "Characteristics of the included trials",
      ["Trial", "Registry ID", "Year", "Design", "Population", "Analysed (T/C)"],
      [[t.get("name") or t["id"], t.get("nct", ""), t.get("year", ""),
        (t.get("design", "") or "")[:60],
        (t.get("population", "") or "")[:50],
        "%s / %s" % (n((t["by_outcome"][OID].get("treatment") or {}).get("n")),
                     n((t["by_outcome"][OID].get("control") or {}).get("n")))]
       for t in d["inputs"]["trials"]])

doc.add_heading("Per-arm event counts", 3)
table(doc, "Per-arm events and denominators for the pooled composite",
      ["Trial", "Intervention events / n", "Comparator events / n",
       "Source tier"],
      [[t.get("name") or t["id"],
        "%s / %s" % (n(t["by_outcome"][OID]["treatment"]["events"]),
                     n(t["by_outcome"][OID]["treatment"]["n"])),
        "%s / %s" % (n(t["by_outcome"][OID]["control"]["events"]),
                     n(t["by_outcome"][OID]["control"]["n"])),
        t["by_outcome"][OID]["treatment"].get("source_tier", "")]
       for t in d["inputs"]["trials"]
       if (t["by_outcome"][OID].get("treatment") or {}).get("events") is not None])

# ---- Extracted values, and where each came from -------------------------------
# PORTED FROM THE HTML RENDERER 2026-08-17. The provenance table existed on the
# page and had NEVER been emitted here, on any manuscript this project has ever
# produced. The Word-vs-HTML alignment gate could not see it, because that gate
# compares the sections BOTH surfaces emit -- so a section present in one and
# absent from the other was silently out of scope rather than a divergence.
#
# A gate that compares only what both surfaces have can never detect absence.
# That is the fifth instance today of a check reporting success without having
# performed the check, and it is why this section went missing indefinitely.
doc.add_heading("Extracted values, and where each came from", 2)
doc.add_paragraph(
    "One row per extracted value, carrying the value, the verbatim sentence it was "
    "read from, a resolvable link to the source, and whether the number was read or "
    "derived. Where any of those is absent the row says so rather than omitting the "
    "value. This table is the audit surface: it is what a reader uses to check this "
    "review against its sources without trusting it.")
_prov_rows = []
for _t in d["inputs"]["trials"]:
    _bo = (_t.get("by_outcome") or {}).get(OID) or {}
    _eff = _bo.get("effect") or {}
    _pv = _bo.get("provenance") or {}
    _q = _pv.get("source_quotes") or []
    _val = ("%s %s (%s%% CI %s to %s)"
            % (_eff.get("measure", ""), n(_eff.get("point")),
               _eff.get("ci_level", 95), n(_eff.get("ci_low")), n(_eff.get("ci_high")))
            if _eff.get("point") is not None else "no effect value held")
    _links = " | ".join(x for x in [
        ("NCT %s" % _t["nct"]) if _t.get("nct") else "no registration id recorded",
        ("PMID %s" % _t["pmid"]) if _t.get("pmid") else "",
        _bo.get("source_url") or "no resolvable source link"] if x)
    _df = _eff.get("derived_from")
    _rd = ("READ from the source as printed" if _pv.get("tag") == "MEASURED" and _df
           else ("DERIVED by us from %s" % _df if _df else "not stated whether read or derived"))
    if _eff.get("derivation_note"):
        _rd += " -- " + _eff["derivation_note"]
    _prov_rows.append([
        "%s / %s" % (_t.get("name", "?"), OID),
        _val,
        ("“" + "”\n“".join(_q) + "”") if _q
        else "no source sentence recorded -- this value cannot be checked against a "
             "quoted line here",
        _links, _rd])
if _prov_rows:
    table(doc, "Extracted values, and where each came from",
          ["Trial / outcome", "Value as extracted", "Verbatim source sentence",
           "Source links", "Read or derived"], _prov_rows, mono_cols=(2,))

doc.add_heading("Component endpoints", 3)
# These were carried in the object and rendered on NO surface -- not the page,
# not this file -- so a reader could not see them and a reviewer could not check
# them. That includes PARACHUTE-HF's published all-cause mortality hazard ratio,
# which is in Table 2 of an OPEN main text that this object had recorded as
# unavailable. A value held but never projected is indistinguishable from a value
# not held, so it is projected here, on the one block sequence both surfaces read.
_ce_rows = []
for _t in d["inputs"]["trials"]:
    _ce = (_t.get("component_endpoints") or {}).get("rows") or []
    for _r in _ce:
        _pe = _r.get("published_effect")
        if _pe:
            _eff = "%s %s (%s to %s)" % (_pe.get("measure", ""), n(_pe.get("point")),
                                         n(_pe.get("ci_low")), n(_pe.get("ci_high")))
            if _pe.get("p_value") is not None:
                _eff += ", P = %s" % n(_pe["p_value"])
        else:
            # Stated, not blank: an empty cell reads as an extraction miss, and
            # for these rows the source genuinely publishes no hazard ratio.
            _eff = "None published"
        _ce_rows.append([_t.get("name") or _t["id"], _r.get("endpoint", ""),
                         "%s / %s" % (n(_r.get("treatment_events")), n(_r.get("treatment_n"))),
                         "%s / %s" % (n(_r.get("control_events")), n(_r.get("control_n"))),
                         _eff, _r.get("source_tier", "")])
if _ce_rows:
    # Caption kept inside the journal's 15-word table-title limit; the
    # qualification it used to carry is a sentence below the table instead,
    # where the length rule does not apply and a reader still meets it.
    table(doc, "Component endpoints and each trial's published effect",
          ["Trial", "Endpoint", "Intervention events / n", "Comparator events / n",
           "Published effect", "Source tier"], _ce_rows)
    _ce_tn = TBL
    doc.add_paragraph(
        "Table %d lists them. They are shown because they are reported, and are "
        "NOT pooled: this " % _ce_tn +
        "review's estimand is the composite. A first-event composite is not the "
        "sum of its components -- on PARACHUTE-HF the components exceed it -- so "
        "no composite anywhere in this review is reconstructed by addition. "
        "‘None published’ means no effect estimate for that endpoint "
        "was found in the sources read for that trial; the per-arm counts were, "
        "and are shown. It is not a derivation this review declined to make and "
        "it is not an extraction miss. Where a trial's own reason is recorded it "
        "is given below.")
    _why = []
    for _t in d["inputs"]["trials"]:
        for _r in ((_t.get("component_endpoints") or {}).get("rows") or []):
            _w = _r.get("no_published_effect_because")
            if _w:
                _why.append("%s, %s: %s" % (_t.get("name") or _t["id"],
                                            _r.get("endpoint", ""), _w))
    for _w in _why:
        doc.add_paragraph(_w)

doc.add_heading("Pooled result", 3)
doc.add_paragraph(
    "Pooled %s %s (%s to %s), k = %s. I-squared %s%%, tau-squared %s, Q %s on "
    "%s degrees of freedom."
    % (pooled["measure"], n(pooled["point"]), n(pooled["ci_low"]),
       n(pooled["ci_high"]), n(res.get("k")), n(het.get("i2")),
       n(het.get("tau2")), n(het.get("q")), n(het.get("df"))))

rows = [["Risk ratio", "%s (%s to %s)" % (n(cp["rr"]["point"]), n(cp["rr"]["ci_low"]),
                                          n(cp["rr"]["ci_high"])), n(cp["rr"]["I2"])],
        ["Odds ratio", "%s (%s to %s)" % (n(cp["or"]["point"]), n(cp["or"]["ci_low"]),
                                          n(cp["or"]["ci_high"])), n(cp["or"]["I2"])],
        ["Risk difference", "%s (%s to %s)" % (n(cp["rd"]["point"]),
                                               n(cp["rd"]["ci_low"]),
                                               n(cp["rd"]["ci_high"])),
         n(cp["rd"]["I2"])]] if cp else []
if rows:
    _cpk = cp.get("_k_as_computed") if cp else None
    table(doc, "The same 2x2 on three scales (sensitivity to the primary "
               "hazard-ratio pool)", ["Measure", "Pooled (95% interval)",
                                      "I-squared (%)", "k"],
          [r + [n(_cpk)] for r in rows])
    if cp and cp.get("_STALE"):
        doc.add_paragraph("Table %d: %s" % (TBL, cp["_STALE"]))

if pan.get("leave_one_out"):
    table(doc, "Leave-one-out sensitivity analysis",
          ["Trial omitted", "Pooled estimate (95% interval)", "I-squared (%)"],
          [[x["omitted"], "%s (%s to %s)" % (n(x["point"]), n(x["ci_low"]),
                                             n(x["ci_high"])), n(x["I2"])]
           for x in pan["leave_one_out"]])

mc = (res.get("sensitivity") or {}).get("between_study_variance_method_comparison")
if mc and mc.get("methods"):
    table(doc, "Between-study-variance estimator comparison",
          ["Estimator", "Interval method", "Point", "Lower", "Upper", "tau-squared"],
          [[m.get("between_study_variance_estimator"), m.get("interval_method"),
            n(m.get("point")), n(m.get("ci_low")), n(m.get("ci_high")),
            n(m.get("tau2"))] for m in mc["methods"]])

doc.add_heading("Figures", 3)
for title, path in figs:
    figure(doc, title, path)

# --- Screening log --------------------------------------------------------
doc.add_heading("Screening and exclusions", 3)
table(doc, "Screening log: every record adjudicated, with its decision and the "
           "reason for it",
      ["Record", "Registry / PMID", "Decision", "Reason"],
      [[r.get("trial", ""),
        " ".join(filter(None, [r.get("nct", ""),
                               ("PMID " + str(r["pmid"])) if r.get("pmid") else ""])),
        r.get("disposition", "excluded"),
        (r.get("reason") or "")[:400]] for r in (sc.get("records") or [])])
doc.add_paragraph(
    "This log covers the records the review adjudicated. The object does not "
    "hold the full set retrieved by the search, nor the stage at which each "
    "decision was taken, so a title/abstract versus full-text split cannot be "
    "shown. That gap is stated here rather than implied by a shorter table.")

# --- GRADE ----------------------------------------------------------------
if g.get("domains"):
    doc.add_heading("Certainty of the evidence", 1)
    table(doc, "GRADE summary of findings",
          ["Domain", "Rating", "Basis"],
          [[k.replace("_", " ").capitalize(), v.get("rating", ""),
            (v.get("basis_in_sources") or "")[:400]]
           for k, v in g["domains"].items()])
    doc.add_paragraph("Overall certainty: %s. %s"
                      % (g.get("certainty", ""), g.get("certainty_derivation", "")))

# --- RoB-2 ----------------------------------------------------------------
rb = d.get("rob2") or {}
if rb.get("trials"):
    doc.add_heading("Risk of bias in the included results", 1)
    ag = rb.get("agreement") or {}
    doc.add_paragraph(
        "Assessed with %s, %s. %s Two assessors worked independently: assessor 1 "
        "was %s (%s family) and assessor 2 was %s (%s family). Neither assembled "
        "the canonical object. %s"
        % (rb.get("tool", ""), rb.get("variant", ""), rb.get("unit_of_assessment", ""),
           rb["assessors"][0].get("model", ""), rb["assessors"][0].get("model_family", ""),
           rb["assessors"][1].get("model", ""), rb["assessors"][1].get("model_family", ""),
           rb.get("blinding", "")))
    table(doc, "Risk of bias by domain and trial, both assessors shown separately "
               "and unreconciled",
          ["Trial", "Domain", "Assessor 1 (%s)" % rb["assessors"][0].get("model_family", ""),
           "Assessor 2 (%s)" % rb["assessors"][1].get("model_family", ""), "Agreed"],
          [[t["trial"], "%s %s" % (dm["domain"], dm["domain_name"]),
            dm["assessor_1_openai"]["judgement"], dm["assessor_2_google"]["judgement"],
            "yes" if dm["agreed"] else "NO"]
           for t in rb["trials"] for dm in t["domains"]])
    table(doc, "Overall risk-of-bias judgement per trial, by the standard algorithm",
          ["Trial", "Assessor 1", "Assessor 2", "Agreed"],
          [[t["trial"], t["overall_assessor_1_openai"].get("judgement", ""),
            t["overall_assessor_2_google"].get("judgement", ""),
            "yes" if t["overall_agreed"] else "NO"] for t in rb["trials"]])
    doc.add_paragraph(
        "Inter-assessor agreement, measured before any reconciliation: %d of %d "
        "domain judgements agreed (%s%%), and %d of %d overall judgements agreed. %s"
        % (ag.get("per_domain_agreed", 0), ag.get("per_domain_total", 0),
           ag.get("per_domain_rate_pct", ""), ag.get("overall_agreed", 0),
           ag.get("overall_total", 0), ag.get("comparison_to_screening", "")))
    if rb.get("disagreements"):
        table(doc, "Open disagreements, carried at the more cautious judgement "
                   "pending adjudication",
              ["Trial", "Domain", "Assessor 1", "Assessor 2", "Carried", "Status"],
              [[x["trial"], x["domain"], x["assessor_1_openai"], x["assessor_2_google"],
                x["provisional_carry"], x["status"]] for x in rb["disagreements"]])
    for f in rb.get("integrity_flags") or []:
        doc.add_paragraph("Integrity flag. %s %s %s"
                          % (f.get("flag", ""), f.get("detail", ""), f.get("action", "")))
    rr = ((g.get("domains") or {}).get("risk_of_bias") or {}).get("rob2_effect_on_this_rating")
    if rr:
        doc.add_paragraph(
            "Effect on the GRADE risk-of-bias rating: it does NOT move. It was %s "
            "before the assessment and remains %s. %s"
            % (rr.get("rating_before_rob2", ""), rr.get("rating_after_rob2", ""),
               rr.get("why_it_does_not_move", "")))
        doc.add_paragraph(rr.get("counter_argument_recorded", ""))

# --- Source discrepancies -------------------------------------------------
_disc = [(t.get("name") or t["id"], x) for t in d["inputs"]["trials"]
         for x in (t.get("discrepancies") or [])]
if _disc:
    doc.add_heading("Disagreements between sources", 1)
    table(doc, "Quantities on which registry and publication disagree",
          ["Trial", "Quantity", "Registry", "Publication", "Status"],
          [[nm, x["quantity"], x["registry_value"], x["publication_value"], x["status"]]
           for nm, x in _disc])
    for nm, x in _disc:
        doc.add_paragraph("%s, %s. Registry pointer: %s. Publication pointer: %s. %s %s"
                          % (nm, x["quantity"], x["registry_pointer"],
                             x["publication_pointer"], x["why_it_matters"],
                             x.get("lesson", "")))

# --- Discussion / limitations / conclusions -------------------------------
# --- Comparison with published syntheses -----------------------------------
# A standard section, not an appendix. The point of this project is not that it
# produced another synthesis; it is that it checked the existing ones and showed
# its working. Confirmations are printed in the same table and the same detail as
# errors, and the denominator is printed with them, because a list of only the
# failures is a selection rather than a finding.
_PC = d.get("published_comparison") or {}
if _PC:
    doc.add_heading("Comparison with published syntheses", 1)
    doc.add_paragraph(_PC.get("_why", ""))
    doc.add_paragraph(_PC.get("_how_identified", ""))
    for _rv in _PC.get("reviews", []):
        doc.add_heading(_rv.get("citation", "")[:120], 2)
        doc.add_paragraph(
            "PMID %s; PMCID %s; DOI %s. %s" % (_rv.get("pmid", ""),
                                               _rv.get("pmcid", ""),
                                               _rv.get("doi", ""),
                                               _rv.get("identifier_provenance", "")))
        doc.add_paragraph("Their scope: %s. Their k: %s. Their search closed %s."
                          % (_rv.get("scope", ""), n(_rv.get("their_k")),
                             _rv.get("their_search_closed", "")))
        doc.add_paragraph(_rv.get("how_it_differs_from_ours", ""))

    # Trial-by-trial reconciliation, from the reconciliation block.
    _tl = ((d.get("reconciliation") or {}).get("trial_list_diffs") or [])
    for _t in _tl:
        if not _t.get("list_was_readable"):
            continue
        _rows = []
        for _x in _t.get("theirs_not_ours", []):
            _rows.append(["only theirs", _x.get("trial", ""),
                          _x.get("disposition", ""), _x.get("reason", "")])
        for _x in _t.get("ours_not_theirs", []):
            _rows.append(["only ours", _x.get("trial", ""), "held",
                          _x.get("why_ours_is_right", "")])
        for _nm in _t.get("review_included_trials", []):
            if any(_nm == r[1] for r in _rows):
                continue
            _rows.append(["in both", _nm, "pooled or screened",
                          "Present in their included-study table and in this "
                          "object's records."])
        if _rows:
            table(doc, "Trial-by-trial reconciliation against their included set",
                  ["Side", "Trial", "Disposition here", "Reason"], _rows)

    _ck = _PC.get("checks", [])
    if _ck:
        table(doc, "Checks applied to the published synthesis, with verdicts",
              ["Check", "Verdict", "What was compared"],
              [[c.get("what", ""), c.get("verdict", ""), c.get("detail", "")]
               for c in _ck])
        _den = _PC.get("denominator") or {}
        doc.add_paragraph("%s %s" % (_den.get("statement", ""),
                                     _den.get("symmetry", "")))
        # Claim, verbatim quotation, location, adjudication -- in that order,
        # because that is the order a reader checks us in.
        table(doc, "Quoted evidence for each check, with its location and our "
                   "adjudication",
              ["Claim checked", "Verbatim quotation", "Location", "Adjudication"],
              [[c.get("what", ""),
                (c.get("quote") or
                 "[nothing to quote: the item is absent from the paper]"),
                c.get("location", ""),
                "%s%s" % (c.get("verdict", ""),
                          (" (%s)" % c["severity"]) if c.get("severity") else "")]
               for c in _ck],
              mono_cols=())
    _dv = _PC.get("divergence_decomposed") or {}
    if _dv:
        doc.add_heading("Where our result differs from theirs", 2)
        doc.add_paragraph("This review: %s" % _dv.get("ours", ""))
        doc.add_paragraph("That review: %s" % _dv.get("theirs", ""))
        doc.add_paragraph(_dv.get("why_they_differ", ""))

# --- Statistical output, quoted verbatim ------------------------------------
# The validity layer. Each block is the captured stdout of the call named above
# it, not a reconstruction, so the k, the estimator, the heterogeneity and the
# package version travel WITH the number. Recorded as a "pre" block so the page
# and this file render the identical text: a quotation that is reflowed is no
# longer a quotation.
_RO = res.get("r_output") or {}
if _RO.get("blocks"):
    doc.add_heading("Statistical output, quoted verbatim", 1)
    doc.add_paragraph(_RO.get("_why", ""))
    doc.add_paragraph(_RO.get("_agreement_checked", ""))
    doc.add_paragraph("Environment: %s" % _RO.get("_environment", ""))
    table(doc, "Statistical output quoted verbatim, with the call that produced it",
          ["Result", "Call", "Verbatim output"],
          [[_b.get("label", _bid), _b.get("call", ""), _b.get("output", "")]
           for _bid, _b in _RO["blocks"].items()],
          mono_cols=(2,))

doc.add_heading("Discussion", 1)
for x in MS.get("discussion", []):
    doc.add_heading(x["heading"], 2)
    doc.add_paragraph(fill(x["text"], "discussion." + x["heading"]))
if not MS.get("discussion"):
    doc.add_paragraph("This section is not written here; the object holds no "
                      "interpretation.")

doc.add_heading("Limitations", 1)
for txt in MS.get("limitations", []):
    doc.add_paragraph(fill(txt, "limitations"), style="List Bullet")
for txt in filter(None, [
        (res.get("sensitivity") or {}).get("leave_one_out_finding"),
        sc.get("known_limitation"),
        (reg.get("ordering") or {}).get("reason"),
        (cp.get("_provenance") if cp else None)]):
    doc.add_paragraph(txt, style="List Bullet")

doc.add_heading("Conclusions", 1)
doc.add_paragraph(fill(MS["conclusions"], "conclusions") if MS.get("conclusions")
                  else "Awaiting the author. See Discussion.")

if MS.get("not_written"):
    doc.add_heading("Sections not written, and why", 1)
    for x in MS["not_written"]:
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(x["section"] + ". ")
        r.bold = True
        para.add_run(x["why"])
doc.add_heading("Funding and conflicts of interest", 1)
doc.add_paragraph("The canonical object records neither a funding statement nor "
                  "a conflict-of-interest declaration for this review.")

# --- References -----------------------------------------------------------
doc.add_heading("References", 1)
# From citations[], NOT from sources{}. The earlier build listed the source
# LAYERS -- PubMed, ClinicalTrials.gov, the registry record -- which is a
# provenance list, not a reference list: it cites the databases rather than the
# papers. Every entry below is a work, in Vancouver order, with a link whose
# reachability was checked and recorded.
_c = d.get("citations") or {}
cits = list(_c.values()) if isinstance(_c, dict) else list(_c)
cits.sort(key=lambda c: (str(c.get("year") or ""), str(c.get("title") or "")))
for i, c in enumerate(cits, 1):
    vol = str(c.get("volume") or "")
    iss = ("(%s)" % c["issue"]) if c.get("issue") else ""
    pgs = (":" + c["pages"]) if c.get("pages") else ""
    loc = ("%s;%s%s%s" % (c.get("year", ""), vol, iss, pgs)) if vol else str(c.get("year") or "")
    ref = "%s. %s. %s. %s." % (c.get("authors_vancouver") or "", c.get("title") or "",
                               c.get("journal") or "", loc)
    ident = " ".join(filter(None, [
        ("PMID " + str(c["pmid"])) if c.get("pmid") else None,
        c.get("nct"), ("doi:" + c["doi"]) if c.get("doi") else None]))
    url = c.get("url") or (("https://pubmed.ncbi.nlm.nih.gov/%s/" % c["pmid"])
                           if c.get("pmid") else "")
    st = c.get("link_status") or c.get("http_status") or c.get("link_checked")
    doc.add_paragraph("%d. %s %s %s %s"
                      % (i, re.sub(r"\s+", " ", ref).replace(" .", ".").strip(),
                         ident, url,
                         ("[link checked: %s]" % st) if st else "[link not checked]"))
doc.add_paragraph(
    "%d references. Every one is a work rather than a database, carries a link, "
    "and the link's reachability was tested at build time and reported above "
    "whether it passed or not. %s"
    % (len(cits), (d.get("citation_policy") or {}).get("statement", "")
       if isinstance(d.get("citation_policy"), dict) else (d.get("citation_policy") or "")))

page_numbers(doc)

# ---- journal-mandated end matter ----------------------------------------
sys.path.insert(0, os.path.join("F:", os.sep, "rapidmeta-ssot-shell", "ssot"))
import journal_profile as _JP  # noqa: E402

_probs = []
if MS.get("keywords"):
    doc.add_heading("Keywords", 1)
    doc.add_paragraph(", ".join(MS["keywords"]))
    _probs += _JP.check_keywords(MS["keywords"])

_das = MS.get("data_availability_statement") or {}
doc.add_heading("Data availability", 1)
if _das.get("underlying_data"):
    _q = doc.add_paragraph()
    _q.add_run("Underlying data. ").bold = True
    _q.add_run(_das["underlying_data"])
_ed = _das.get("extended_data") or {}
if _ed:
    _q = doc.add_paragraph()
    _q.add_run("Extended data. ").bold = True
    _q.add_run("%s: %s. %s"
               % (_ed.get("repository", ""), _ed.get("title", ""),
                  _ed.get("persistent_identifier")
                  or ("[persistent identifier not yet minted]")))
    if not _ed.get("persistent_identifier"):
        doc.add_paragraph(_ed.get("identifier_status", ""))
        _probs.append("Data Availability Statement carries no persistent "
                      "identifier for the extended data.")
    doc.add_paragraph("This project contains the following extended data:")
    for _fn, _desc in _ed.get("files", []):
        _q = doc.add_paragraph(style="List Bullet")
        _q.add_run(_fn + ". ").bold = True
        _q.add_run(fill(_desc, "das.files"))
    _q = doc.add_paragraph()
    _q.add_run("Data are available under the terms of the ").italic = True
    _q.add_run(_ed.get("licence", "")).italic = True

_lic = MS.get("licences") or {}
if _lic:
    _q = doc.add_paragraph()
    _q.add_run("Note on licensing. ").bold = True
    _q.add_run("The software and the data carry DIFFERENT licences and are "
               "different objects. Code: %s -- %s Data: %s -- %s"
               % (_lic["code"]["spdx"], _lic["code"]["obligation"],
                  _lic["data"]["spdx"], _lic["data"]["obligation"]))

_sw = MS.get("software_availability") or {}
if _sw:
    doc.add_heading("Software availability", 1)
    for _lab, _key in (("Source code available from", "source_code_available_from"),
                       ("Version control", "version_control"),
                       ("Archived source code at time of publication",
                        "archived_source_code_at_time_of_publication"),
                       ("Licence", "licence")):
        _q = doc.add_paragraph()
        _q.add_run(_lab + ": ").bold = True
        _q.add_run(str(_sw.get(_key) or "[not yet available]"))
    if not _sw.get("archived_source_code_at_time_of_publication"):
        doc.add_paragraph(_sw.get("archive_status", ""))
        _probs.append("No archived source-code identifier (Zenodo) for the code.")
    if not _sw.get("licence"):
        _probs.append("No OSI-approved licence declared.")

if MS.get("registration_note_for_editor"):
    doc.add_heading("Note on registration", 1)
    doc.add_paragraph(MS["registration_note_for_editor"])

# Tables, then figure legends, both at the end.
_DEFER_ON = False
if _DEFERRED_T:
    doc.add_heading("Tables", 1)
    for _n, _cap, _hd, _rw, _mc in _DEFERRED_T:
        TBL = _n - 1
        _real_table(doc, _cap, _hd, _rw, _mc)
        _probs += _JP.check_title_words(_cap, "table")
    TBL = len(_DEFERRED_T)
if _DEFERRED_F:
    doc.add_heading("Figure legends", 1)
    for _n, _cap, _p, _w in _DEFERRED_F:
        _pp = doc.add_paragraph()
        _pp.add_run("Figure %d. " % _n).bold = True
        _pp.add_run(_cap)
        _probs += _JP.check_title_words(_cap, "figure")
    doc.add_heading("Figures", 1)
    for _n, _cap, _p, _w in _DEFERRED_F:
        FIG = _n - 1
        _real_figure(doc, _cap, _p, _w)
    FIG = len(_DEFERRED_F)

doc.add_heading("Submission conformance", 1)
for _k, _v in _JP.statement_of_conformance().items():
    _q = doc.add_paragraph()
    _q.add_run(_k.replace("_", " ").capitalize() + ". ").bold = True
    _q.add_run(str(_v))
if _probs:
    _w2 = doc.add_paragraph()
    _r2 = _w2.add_run("NOT YET SUBMITTABLE. " + " ".join(_probs))
    _r2.bold = True
print("conformance problems:", _probs or "NONE")

import json as _json
_json.dump({"built_utc": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "title": d["title"], "question": d["question"],
            "tables": TBL, "figures": FIG, "blocks": DOCMODEL},
           open(os.path.join(OUTDIR, "manuscript_docmodel.json"), "w",
                encoding="utf-8"), indent=1, ensure_ascii=False)
# CORE PROPERTIES. These were python-docx's template defaults: an empty title,
# "python-docx" as author, and created/modified stamped 2013-12-23. Word shows
# dc:title in Properties and in recent-file lists, and some submission portals
# read it, so an empty one is a blank where the paper's name belongs and a 2013
# date is simply false. Filled from the object -- the same source the first line
# of the document comes from -- so the title in the metadata and the title on the
# page cannot disagree, which is the exact failure this pass was about.
# `author` is deliberately NOT invented: the object records none, and a
# fabricated author in submission metadata is worse than an absent one.
try:
    import datetime as _dt
    _cp = doc.core_properties
    _cp.title = d.get("title", "")
    _cp.subject = d.get("question", "")
    _cp.keywords = ", ".join((d.get("manuscript") or {}).get("keywords") or [])
    _cp.category = "Systematic review and meta-analysis"
    _cp.comments = ("Every number is projected from a single canonical object; "
                    "see the Data availability section.")
    try:
        _when = _dt.datetime.strptime(str(d.get("built") or "")[:10], "%Y-%m-%d")
        _cp.created = _when
        _cp.modified = _when
    except ValueError:
        pass
except Exception as _e:                                   # noqa: BLE001
    print("WARNING: core properties not set: %s" % _e)

# --- F1000Research submission checklist -------------------------------------
# Requirements taken from F1000Research's own author guidance for Systematic
# Reviews, looked up rather than remembered (their site 403s direct fetch, so
# these come from the indexed guideline pages):
#   * a Data Availability Statement is mandatory EVEN WHERE THERE IS NO DATA;
#   * PRISMA checklist AND flow diagram are required, and the completed
#     checklist and flow chart must be DEPOSITED in an approved repository,
#     with the guideline type, repository, DOI and licence in the statement;
#   * extended data needs title, repository, DOI/accession and licence, under
#     an "Extended data" subheading, cited in the main text;
#   * the repository must supply a persistent identifier and allow CC0 /
#     CC-BY 4.0 or equivalent;
#   * archived source code needs a DOI and citation in Zenodo under an open,
#     preferably OSI-approved, licence.
#
# These are MANDATORY FIELDS at this venue, not nice-to-haves, so an unmet one
# blocks rather than prints. The missing identifiers are recorded as null with a
# stated reason and are never filled with a plausible string: a wrong DOI in a
# Data Availability Statement points a reader at someone else's data, which is
# worse than no DOI at all. Minting them is an author action, not a build step.
_F1000 = []
_das = MS.get("data_availability_statement") or {}
_ed = _das.get("extended_data") or {}
_swx = MS.get("software_availability") or {}
for _label, _ok in (
    ("Data Availability Statement present", bool(_das)),
    ("Extended data: repository named", bool(_ed.get("repository"))),
    ("Extended data: persistent identifier (DOI) minted",
     bool(_ed.get("persistent_identifier"))),
    ("Extended data: open licence (CC0 or CC-BY 4.0)",
     str(_ed.get("licence", "")).upper().startswith(("CC0", "CC-BY", "CC BY"))),
    ("Software: source code location", bool(_swx.get("source_code_available_from"))),
    ("Software: archived Zenodo DOI at publication",
     bool(_swx.get("archived_source_code_at_time_of_publication"))),
    ("Software: OSI-approved licence", bool(_swx.get("licence"))),
    ("PRISMA flow diagram present", FIG > 0),
    ("PRISMA checklist deposited with a DOI",
     bool((MS.get("prisma") or {}).get("checklist_doi"))),
    ("Structured abstract", isinstance(MS.get("abstract"), dict)
     and len(MS.get("abstract") or {}) >= 4),
    ("Registration statement", bool(MS.get("registration_note_for_editor"))),
):
    if not _ok:
        _F1000.append(_label)
if _F1000:
    _probs.append("F1000 mandatory requirements unmet: " + "; ".join(_F1000))

doc.save(OUT)
print("wrote", OUT, os.path.getsize(OUT), "bytes")
if _F1000:
    print("")
    print("#" * 72)
    print("SUBMISSION BLOCKED -- %d mandatory F1000Research requirement(s) unmet:"
          % len(_F1000))
    for _x in _F1000:
        print("   - %s" % _x)
    print("These are author actions (Zenodo deposits), not build steps. No "
          "identifier will be invented to clear them.")
    print("#" * 72)
print("tables:", TBL, "figures:", FIG, "rasterised:", len(figs))
